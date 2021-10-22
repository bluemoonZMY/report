from docx import Document
from docx.shared import Cm,Inches, Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #对齐
from docx.enum.text import WD_LINE_SPACING
from add_float_picture import add_float_picture
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
with open("./GGJ02.2020.MY03.01.514.json",encoding="utf-8") as ff:
	datasourcef=json.load(ff)  
def output():
	r=7
	c=4
	document = Document('./output.docx')
	table=document.add_table(rows=r,cols=c,style='Table Grid')
	#合并之前设置table行高
	for rr in range(4):
			table.rows[rr].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
			table.rows[rr].height=Cm(1.4)
	#指定table列宽
	for rr in range(7):
			table.cell(rr,0).width=Cm(2.25)
			table.cell(rr,1).width=Cm(6.25)
			table.cell(rr,2).width=Cm(2)
			table.cell(rr,3).width=Cm(5)
	A = table.cell(2, 2).merge(table.cell(3, 2))
	B = table.cell(5, 1).merge(table.cell(5, 3))
	C = table.cell(2, 3).merge(table.cell(3, 3))
	D = table.cell(6, 1).merge(table.cell(6, 3))
	table.cell(0,0).text='产品名称'
	table.cell(1,0).text='制造商'
	table.cell(2,0).text='委托方'
	table.cell(3,0).text='检验类型'
	table.cell(4,0).text='检验\n依据'
	table.cell(5,0).text='检\n \n验\n \n结\n \n论'
	table.cell(6,0).text='备\n注'
	table.cell(0,2).text='产品型号'
	table.cell(1,2).text='商  标'
	A.text='检验\n日期'
	table.cell(4,2).text='检验\n项目'
	##table.cell(0,3).text=datasource.['签章页'].['产品型号']
	D.text='（1）任务来源：委    托   （5）试验结果：见附录D\n（2）试验对象：见附录A   （6）参试人员：见附录E\n（3）试验依据：见附录B   （7）试验照片：见附录F\n（4）试验条件：见附录C'
	#####添加一行空白
	maincontent=document.add_paragraph() #主要内容
	maincontent.alignment=WD_ALIGN_PARAGRAPH.LEFT
	mainContentrun= maincontent.add_run(' ')
	mainContentrun.font.size=Pt(10.5)
	mainContentrun.font.name=u'宋体'
	mainContentrun._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	maincontent.paragraph_format.line_spacing=1
	maincontent.paragraph_format.space_after=Pt(0)
	maincontent.paragraph_format.space_after=Pt(0)
	#添加插章表(没有实线)
	table1=document.add_table(rows=1,cols=6)
	#行高固定2Cm
	table1.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
	table1.rows[0].height=Cm(2)
	#设置文字样式
	#表格文字先统一居中
	table.style.font.size=Pt(12)
	for i in range(r):
		for j in range(c):
				table.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
	table.style.font.name=u'宋体'
	table.style._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	table.style.paragraph_format.line_spacing=1
	table.style.paragraph_format.space_after=Pt(0)
	table.style.paragraph_format.space_after=Pt(0)
	#居中对齐
	table.cell(0,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(0,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(0,2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(0,3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(1,2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(1,3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(2,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(2,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(3,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(3,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(4,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(4,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(4,2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	table.cell(5,0).paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
	A.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	#左端对齐
	C.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
	table.cell(4,3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
	D.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
	#两端对齐
	table.cell(6,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	B.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
	document.save('./output.docx')