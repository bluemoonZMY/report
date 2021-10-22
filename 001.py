from docx import Document
import json
from docx.shared import Cm,Inches, Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #对齐
from docx.enum.text import WD_LINE_SPACING
from add_float_picture import add_float_picture
import os
#os.chdir(os.path.dirname(__file__))
document = Document()
with open("./GGJ02.2020.MY03.01.514.json",encoding="utf-8") as f:
	datasource=json.load(f)           #要打开的源文件名
sections = document.sections
for section in sections:
	section.top_margin = Cm(2.5)
	section.bottom_margin = Cm(2.5)
	section.left_margin = Cm(3.0)
	section.right_margin = Cm(2.5)
	section.page_width = Cm(21)#页面大小与页边距
	section.page_height = Cm(29.7)
print('默认页面的宽度和高度：', section.page_width.cm,section.page_height.cm)
p = document.add_paragraph()
add_float_picture(p, './资质图片.png', width=Inches(7.36), pos_x=Inches(0.454), pos_y=Pt(60))
add_float_picture(p, './印章.png', width=Inches(3.25), pos_x=Inches(2.51), pos_y=Pt(620))
    
coverReportnumber=document.add_paragraph() #报告编号
coverReportnumber.alignment=WD_ALIGN_PARAGRAPH.RIGHT
coverReportNumberrun=coverReportnumber.add_run('\n\n\n\n\n报告编号：'+datasource["报告编号"]+'\n\n\n')
coverReportNumberrun.font.size=Pt(14)
coverReportNumberrun.font.name=u'宋体'
coverReportNumberrun._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')  
coverReportnumber.paragraph_format.line_spacing=1.0
coverReportnumber.paragraph_format.space_after=Pt(0)  #失效
coverReportnumber.paragraph_format.space_before=Pt(0)
def inspectionReport():
	inspectionreport=document.add_paragraph() #检验报告
	inspectionreport.alignment=WD_ALIGN_PARAGRAPH.CENTER
	inspectionReportrun=inspectionreport.add_run('检验报告')
	inspectionReportrun.font.size=Pt(42)
	inspectionReportrun.font.name='Arial'
	inspectionReportrun._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
	inspectionreport.paragraph_format.line_spacing=1.0
	inspectionreport.paragraph_format.space_after=Pt(0)
	inspectionreport.paragraph_format.space_before=Pt(0)
inspectionReport()
def maincontent():
	
	maincontent=document.add_paragraph() #主要内容
	maincontent.alignment=WD_ALIGN_PARAGRAPH.LEFT
	mainContentrun= maincontent.add_run('\n\n\n\n产品名称：'+datasource["产品名称"]+'\n产品型号：'+datasource["产品型号"]+'\n委托方：'+datasource["委托方"]+'\n检验类型：'+datasource["检验类型"]+'\n\n\n\n')
	mainContentrun.font.size=Pt(16)
	mainContentrun.font.name='Arial'
	mainContentrun._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	maincontent.paragraph_format.line_spacing=1.5
	maincontent.paragraph_format.space_after=Pt(0)
	maincontent.paragraph_format.space_after=Pt(0)
	maincontent.paragraph_format.left_indent=Cm(3.8)
maincontent()
downcontent=document.add_paragraph() #下方内容
downcontent.alignment=WD_ALIGN_PARAGRAPH.CENTER
downContentrun=downcontent.add_run('中机科（北京）车辆检测工程研究院有限公司\n国家工程机械质量监督检验中心\n'+datasource["检验时间"]+'二零二一年四月\n')
downContentrun.font.size=Pt(16)
downContentrun.font.name='Arial'
downContentrun._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
downcontent.paragraph_format.line_spacing=1.5
downcontent.paragraph_format.space_after=Pt(0)
downcontent.paragraph_format.space_after=Pt(0)
document.add_page_break() #分页符
flyleaf1=document.add_paragraph()
flyleaf1.alignment=WD_ALIGN_PARAGRAPH.CENTER
flyleaf1run=flyleaf1.add_run('注意事项\n')
flyleaf1run.font.size=Pt(16)
flyleaf1run.font.name='Arial'
flyleaf1run._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
flyleaf1run.bold=True
flyleaf1.paragraph_format.line_spacing=1
flyleaf1.paragraph_format.space_after=Pt(0)
flyleaf1.paragraph_format.space_after=Pt(0)
	 
flyleaf2=document.add_paragraph()
flyleaf2.alignment=WD_ALIGN_PARAGRAPH.LEFT
flyleaf2run=flyleaf2.add_run('1.报告无“检验报告专用章”或型式检验检测机构公章无效；报告无加盖骑缝章无效。\n2.复制报告未重新加盖“检验报告专用章”或型式试验检验检测机构公章无效；复制报告未重新加盖骑缝章无效。\n3.报告无主检、审核、签发人签字无效。\n4.报告涂改无效。\n5.报告是对设备型式的确认，对样品本身的合格与否负责，且仅对符合送样样品的产品有效。')
flyleaf2run.font.size=Pt(14)
flyleaf2run.font.name='Arial'
flyleaf2run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
flyleaf2.paragraph_format.line_spacing=1.5
flyleaf2.paragraph_format.space_after=Pt(0)
flyleaf2.paragraph_format.space_after=Pt(0)
flyleaftable=document.add_table(rows=1,cols=1)
flyleaftable.cell(0,0).width=Cm(15.5)
flyleaftable.cell(0,0).height=Cm(0.04) ##

flyleaf3=document.add_paragraph()
flyleaf3.alignment=WD_ALIGN_PARAGRAPH.CENTER
flyleaf3run=flyleaf3.add_run('报告附加说明')
flyleaf3run.font.size=Pt(16)
flyleaf3run.font.name='Arial'
flyleaf3run._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
flyleaf3run.bold=True
flyleaf3.paragraph_format.line_spacing=1
flyleaf3.paragraph_format.space_after=Pt(0)
flyleaf3.paragraph_format.space_after=Pt(0)
flyleaf4=document.add_paragraph()
flyleaf4.alignment=WD_ALIGN_PARAGRAPH.LEFT
def flyleaf40(part1,part2,part3):
	flyleaf401run=flyleaf4.add_run(part1)
	flyleaf401run.font.size=Pt(14)
	flyleaf401run.font.name=u'宋体'
	flyleaf401run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	flyleaf402run=flyleaf4.add_run(part2)
	flyleaf402run.font.color.rgb= RGBColor(255,255,255)
	flyleaf402run.font.size=Pt(14)
	flyleaf402run.font.name=u'宋体'
	flyleaf402run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	flyleaf403run=flyleaf4.add_run(part3)
	flyleaf403run.font.size=Pt(14)
	flyleaf403run.font.name=u'宋体'
	flyleaf403run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
flyleaf40('1.','—','委托方地址：'+datasource["报告附加说明"]["委托方地址"]+'\n')
flyleaf40('2.','—','委托方电话：'+datasource["报告附加说明"]["委托方电话"]+'\n')
flyleaf40('3.','—','委托方法定代表人：'+datasource["报告附加说明"]["委托方法定代表人"]+'\n')
flyleaf40('4.','—','制造商地址：'+datasource["报告附加说明"]["制造商地址"]+'\n')
flyleaf40('5.','—','制造商电话：'+datasource["报告附加说明"]["制造商电话"]+'\n')
flyleaf40('6.','—','制造商法定代表人：'+datasource["报告附加说明"]["制造商法定代表人"]+'\n')
flyleaf40('7.','—','性能试验与安全要求验证用样机接受日期：'+datasource["报告附加说明"]["性能试验与安全要求验证用样机接受日期"]+'\n')
flyleaf40('8.','—','可靠性试验用样机接受日期'+datasource["报告附加说明"]["可靠性试验用样机接受日期"]+'\n')
flyleaf40('9.','—','试验项目有无外包：'+datasource["报告附加说明"]["试验项目有无外包"]+'\n')
flyleaf40('10.','—','检验检测机构地址：'+datasource["报告附加说明"]["检验检测机构地址"]+'\n')
flyleaf40('11.','—','检验检测机构电话'+datasource["报告附加说明"]["检验检测机构电话"]+'\n')
flyleaf40('12.','—','投诉电话：'+datasource["报告附加说明"]["投诉电话"]+'\n')
document.add_page_break() #分页符
################################################
def content(indent,part1,part2,part3):
	cont=document.add_paragraph()
	cont.alignment=WD_ALIGN_PARAGRAPH.DISTRIBUTE
	cont.paragraph_format.line_spacing=1
	cont.paragraph_format.space_after=Pt(0)
	cont.paragraph_format.space_after=Pt(0)
	cont.paragraph_format.left_indent=Inches(indent)
	run=cont.add_run(part1)
	run.font.size=Pt(12)
	run.font.name=u'宋体'
	run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	run1=cont.add_run(part2)
	run1.font.size=Pt(12)
	run1.font.name=u'宋体'
	run1._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
	run2=cont.add_run(part3)
	run2.font.size=Pt(12)
	run2.font.name=u'宋体'
	run2._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    	#
content (0,'检验结论','……………………………………………………………………………………','1') 
content (0,'附录A 试验对象','…………………………………………………………………………','2')
content (0,'附录B 试验依据','…………………………………………………………………………','7')
content (0,'附录C 试验条件','…………………………………………………………………………','8')
content (0,'附录D 试验结果','…………………………………………………………………………','9')
content (0,'D1试验前样机验收检查结果','……………………………………………………………','9')
content (0,'D2性能试验结果','…………………………………………………………………………','9')
content (0.25,'D2.1外形尺寸测定结果','……………………………………………………………','9')
content (0.25,'D2.2作业参数测定结果','……………………………………………………………','9')
content (0.25,'D2.3质量测定结果','………………………………………………………………','10')
content (0.25,'D2.4接地比压测定结果','…………………………………………………………','10')
content (0.25,'D2.5最大挖掘力测定结果','………………………………………………………','10')
content (0.25,'D2.6行驶速度试验结果','…………………………………………………………','10')
content (0.25,'D2.7爬坡能力试验结果','…………………………………………………………','10')
content (0.25,'D2.8履带式挖掘机行驶直线性测定结果','…………………………………………','11')
content (0.25,'D2.9作业试验结果','…………………………………………………………………','11')
content (0.25,'D2.10强度试验结果','…………………………………………………………………','11')
content (0.25,'D2.11液压系统试验结果','……………………………………………………………','13')
content (0,'D3安全要求和防护措施的验证结果','……………………………………………………','14')
content (0.25,'D3.1通道检查结果','…………………………………………………………………','14')
content (0.25,'D3.2司机操作位置检查结果','………………………………………………………','16')
content (0.25,'D3.3座椅全检查结果','………………………………………………………………','19')
content (0.25,'D3.4司机的操纵装置和指示装置检查结果','………………………………………','20')
content (0.25,'D3.5转向系统检查结果','……………………………………………………………','24')
content (0.25,'D3.6制动系统检查结果','……………………………………………………………','24')
content (0.25,'D3.7可视性检查结果','………………………………………………………………','26')
content (0.25,'D3.8报警装置和安全标志检查结果','………………………………………………','26')
content (0.25,'D3.9稳定性测试结果','………………………………………………………………','31')
content (0.25,'D3.10噪声测试结果','…………………………………………………………………','31')
content (0.25,'D3.11保护措施及装置检查结果','……………………………………………………','32')
content (0.25,'D3.12救助、捆系、起吊、牵引和运输检查结果','…………………………………','34')
content (0.25,'D3.13电气和电子系统检查结果','……………………………………………………','35')
content (0.25,'D3.14压力系统检查结果','……………………………………………………………','37')
content (0.25,'D3.15燃油箱和液压油箱检查结果','…………………………………………………','38')
content (0.25,'D3.16防火检查结果','…………………………………………………………………','38')
content (0.25,'D3.17维修检查结果','…………………………………………………………………','38')
content (0.25,'D3.18使用信息检查结果','……………………………………………………………','39')
content (0,'D4可靠性试验结果','………………………………………………………………………','39')
content (0,'附录E  参试人员','………………………………………………………………………','40')
content (0,'附录F  试验照片','………………………………………………………………………','41')
document.add_page_break() #分页符
document.save('./output.docx')
#结论页
import conclusion
conclusion.output()


#附录A
import appendixA


    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    

